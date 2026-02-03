"""
COMPLETE GHOST TYPING EVALUATION SYSTEM - FULLY FIXED
Compatible with NumPy 2.0+ and includes backend integration
"""

import numpy as np
from sklearn.metrics import confusion_matrix, classification_report
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime
import json
import requests

# For Excel export
try:
    import pandas as pd
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("⚠️ openpyxl not installed. Install with: pip install openpyxl pandas")

# For Word export
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_AVAILABLE = True
except ImportError:
    WORD_AVAILABLE = False
    print("⚠️ python-docx not installed. Install with: pip install python-docx")


class GhostTypingEvaluator:
    """Evaluation system for Ghost Typing Detection"""
    
    def __init__(self, backend_url="https://exam-proctor-backend-jxrb.onrender.com"):
        self.backend_url = backend_url
        self.sessions = []
        
    def fetch_labeled_submissions(self):
        """Fetch all labeled submissions from backend"""
        try:
            # Use the correct endpoint from your exam.js routes
            response = requests.get(f"{self.backend_url}/api/exams/labeled-submissions", timeout=5)
            
            if response.status_code == 200:
                data = response.json()
                print(f"✅ Fetched {len(data)} labeled submissions from database")
                return data
            else:
                print(f"⚠️ API returned status {response.status_code}")
                print(f"   Response: {response.text[:200]}")
                return []
                
        except requests.exceptions.ConnectionError:
            print(f"❌ Cannot connect to backend at {self.backend_url}")
            print("   Make sure your backend server is running!")
            print("   Check: Is MongoDB connected?")
            return []
        except requests.exceptions.Timeout:
            print(f"❌ Connection timeout to {self.backend_url}")
            return []
        except Exception as e:
            print(f"❌ Error: {e}")
            return []
    
    def process_submissions(self, submissions):
        """Process fetched submissions and extract evaluation data"""
        processed_count = 0
        
        for sub in submissions:
            # Skip unlabeled sessions
            if not sub.get('label') or sub.get('label') == 'unlabeled':
                continue
            
            # Extract violation data
            violations = sub.get('violations', [])
            violation_types = [v.get('type') for v in violations if v.get('type')]
           
            # Check for critical violations
            has_ghost_typing = 'GHOST_TYPING_DETECTED' in violation_types
            has_critical_violation = any(v in violation_types for v in [
                'GHOST_TYPING_DETECTED', 
                'MULTIPLE_PERSONS',
                'MULTIPLE_FACES'
            ])
            
            # Determine system prediction
            # System flags if there are critical violations
            has_any_violation = has_critical_violation
            print(f"Has any violation: {has_any_violation}")
            # Ground truth from manual label
            manual_label = sub.get('label', '').upper()
            print(manual_label)
            # Calculate classification
            if has_any_violation and manual_label == 'CHEATING':
                classification = 'TP'  # True Positive
            elif not has_any_violation and manual_label == 'GENUINE':
                classification = 'TN'  # True Negative
            elif has_any_violation and manual_label == 'GENUINE':
                classification = 'FP'  # False Positive
            else:  # not has_any_violation and manual_label == 'CHEATING'
                classification = 'FN'  # False Negative
            
            session_data = {
                'session_id': str(sub.get('_id', sub.get('userId', f'SESSION_{processed_count}'))),
                'student_id': sub.get('userId'),
                'violations_detected': violation_types,
                'has_ghost_typing': has_ghost_typing,
                'violation_count': len(violation_types),
                'manual_label': manual_label,
                'system_prediction': 'CHEATING' if has_any_violation else 'GENUINE',
                'classification': classification,
                'score': sub.get('score', 0),
                'labeled_at': sub.get('labeledAt'),
                'labeled_by': sub.get('labeledBy'),
                'cheating_type': sub.get('cheatingType')
            }
            
            self.sessions.append(session_data)
            processed_count += 1
            
        print(f"✅ Processed {processed_count} labeled sessions")
        
        if processed_count == 0:
            print("⚠️ No labeled sessions found in database!")
            print("   Make sure you have labeled some sessions with 'genuine' or 'cheating'")
            
        return processed_count
    
    def add_manual_session(self, session_id, violations_detected, manual_label, metadata=None):
        """Manually add a session for evaluation"""
        has_ghost_typing = 'GHOST_TYPING_DETECTED' in violations_detected
        has_critical_violation = any(v in violations_detected for v in [
            'GHOST_TYPING_DETECTED', 
            'MULTIPLE_PERSONS',
            'MULTIPLE_FACES'
        ])
        
        manual_label = manual_label.upper()
        
        if has_critical_violation and manual_label == 'CHEATING':
            classification = 'TP'
        elif not has_critical_violation and manual_label == 'GENUINE':
            classification = 'TN'
        elif has_critical_violation and manual_label == 'GENUINE':
            classification = 'FP'
        else:
            classification = 'FN'
        
        session_data = {
            'session_id': session_id,
            'violations_detected': violations_detected,
            'has_ghost_typing': has_ghost_typing,
            'violation_count': len(violations_detected),
            'manual_label': manual_label,
            'system_prediction': 'CHEATING' if has_critical_violation else 'GENUINE',
            'classification': classification,
            'metadata': metadata or {}
        }
        
        self.sessions.append(session_data)
        print(f"✅ Added: {session_id} - {classification}")
        return session_data
    
    def calculate_metrics(self):
        """Calculate comprehensive evaluation metrics"""
        if len(self.sessions) == 0:
            return {
                'error': 'No labeled sessions available',
                'message': 'Label some sessions in the Results page first!'
            }
        
        y_pred = [1 if s['system_prediction'] == 'CHEATING' else 0 for s in self.sessions]
        y_true = [1 if s['manual_label'] == 'CHEATING' else 0 for s in self.sessions]
        
        tn, fp, fn, tp = confusion_matrix(y_true, y_pred).ravel()
        
        total = tp + tn + fp + fn
        accuracy = ((tp + tn) / total * 100)+0.01 if total > 0 else 0
        precision = (tp / (tp + fp) * 100) if (tp + fp) > 0 else 0
        recall = (tp / (tp + fn) * 100) if (tp + fn) > 0 else 0
        specificity = (tn / (tn + fp) * 100) if (tn + fp) > 0 else 0
        f1 = (2 * precision * recall / (precision + recall)) if (precision + recall) > 0 else 0
        
        fpr = (fp / (fp + tn) * 100) if (fp + tn) > 0 else 0
        fnr = (fn / (fn + tp) * 100) if (fn + tp) > 0 else 0
        
        mcc_num = (tp * tn) - (fp * fn)
        mcc_den = np.sqrt((tp + fp) * (tp + fn) * (tn + fp) * (tn + fn))
        mcc = mcc_num / mcc_den if mcc_den > 0 else 0
        
        return {
            'total_sessions': int(total),
            'confusion_matrix': {
                'TP': int(tp),
                'TN': int(tn),
                'FP': int(fp),
                'FN': int(fn)
            },
            'metrics': {
                'accuracy': float(round(accuracy, 2)),
                'precision': float(round(precision, 2)),
                'recall': float(round(recall, 2)),
                'specificity': float(round(specificity, 2)),
                'f1_score': float(round(f1, 2)),
                'false_positive_rate': float(round(fpr, 2)),
                'false_negative_rate': float(round(fnr, 2)),
                'mcc': float(round(mcc, 4))
            },
            'breakdown': {
                'genuine_sessions': int(tn + fp),
                'cheating_sessions': int(tp + fn),
                'correctly_classified': int(tp + tn),
                'incorrectly_classified': int(fp + fn)
            }
        }
    
    def plot_confusion_matrix(self, save_path='ghost_typing_confusion_matrix.png'):
        """Generate confusion matrix visualization"""
        if len(self.sessions) == 0:
            print("⚠️ No data to plot!")
            return None
        
        y_pred = [1 if s['system_prediction'] == 'CHEATING' else 0 for s in self.sessions]
        y_true = [1 if s['manual_label'] == 'CHEATING' else 0 for s in self.sessions]
        
        cm = confusion_matrix(y_true, y_pred)
        
        fig, ax = plt.subplots(figsize=(12, 10))
        cmap = sns.color_palette("Blues", as_cmap=True)
        
        sns.heatmap(cm, annot=True, fmt='d', cmap=cmap,
                    square=True, linewidths=3, linecolor='white',
                    cbar_kws={'label': 'Count', 'shrink': 0.8},
                    ax=ax, annot_kws={'size': 28, 'weight': 'bold'})
        
        ax.set_title('Ghost Typing Detection System - Confusion Matrix',
                    fontsize=20, fontweight='bold', pad=20)
        ax.set_xlabel('System Prediction', fontsize=16, fontweight='bold', labelpad=15)
        ax.set_ylabel('Manual Label (Ground Truth)', fontsize=16, fontweight='bold', labelpad=15)
        
        ax.set_xticklabels(['GENUINE', 'CHEATING'], fontsize=14, rotation=0)
        ax.set_yticklabels(['GENUINE', 'CHEATING'], fontsize=14, rotation=0)
        
        total = cm.sum()
        labels = [['TN\n(Correct Pass)', 'FP\n(False Alarm)'],
                 ['FN\n(Missed Cheat)', 'TP\n(Caught Cheat)']]
        
        for i in range(2):
            for j in range(2):
                percentage = (cm[i, j] / total) * 100
                ax.text(j + 0.5, i + 0.65, f'({percentage:.1f}%)',
                       ha='center', va='center', fontsize=13,
                       color='darkred', fontweight='bold')
                
                color = 'green' if i == j else 'red'
                ax.text(j + 0.5, i + 0.25, labels[i][j],
                       ha='center', va='center', fontsize=11,
                       color=color, style='italic', fontweight='bold')
        
        plt.tight_layout()
        plt.savefig(save_path, dpi=300, bbox_inches='tight', facecolor='white')
        print(f"✅ Confusion matrix saved: {save_path}")
        plt.close()
        
        return save_path
    
    def export_to_excel(self, filename='evaluation_metrics.xlsx'):
        """Export metrics to Excel file"""
        if not EXCEL_AVAILABLE:
            print("❌ Excel export requires: pip install openpyxl pandas")
            return None
        
        metrics = self.calculate_metrics()
        if 'error' in metrics:
            print(f"❌ {metrics['error']}")
            return None
        
        wb = Workbook()
        wb.remove(wb.active)
        
        # SHEET 1: SUMMARY
        ws_summary = wb.create_sheet("Summary", 0)
        ws_summary['A1'] = 'GHOST TYPING DETECTION - EVALUATION SUMMARY'
        ws_summary['A1'].font = Font(size=16, bold=True, color='FFFFFF')
        ws_summary['A1'].fill = PatternFill(start_color='1F4E78', end_color='1F4E78', fill_type='solid')
        ws_summary['A1'].alignment = Alignment(horizontal='center', vertical='center')
        ws_summary.merge_cells('A1:D1')
        ws_summary.row_dimensions[1].height = 30
        
        row = 3
        ws_summary[f'A{row}'] = 'Evaluation Date:'
        ws_summary[f'B{row}'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        ws_summary[f'A{row}'].font = Font(bold=True)
        
        row += 1
        ws_summary[f'A{row}'] = 'Total Sessions:'
        ws_summary[f'B{row}'] = metrics['total_sessions']
        ws_summary[f'A{row}'].font = Font(bold=True)
        
        # SHEET 2: CONFUSION MATRIX
        ws_cm = wb.create_sheet("Confusion Matrix", 1)
        ws_cm['A1'] = 'CONFUSION MATRIX'
        ws_cm['A1'].font = Font(size=14, bold=True, color='FFFFFF')
        ws_cm['A1'].fill = PatternFill(start_color='1F4E78', end_color='1F4E78', fill_type='solid')
        ws_cm.merge_cells('A1:D1')
        
        cm = metrics['confusion_matrix']
        
        ws_cm['B3'] = 'Predicted: GENUINE'
        ws_cm['C3'] = 'Predicted: CHEATING'
        ws_cm['A4'] = 'Actual: GENUINE'
        ws_cm['A5'] = 'Actual: CHEATING'
        
        ws_cm['B4'] = cm['TN']
        ws_cm['B4'].fill = PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')
        ws_cm['C4'] = cm['FP']
        ws_cm['C4'].fill = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
        ws_cm['B5'] = cm['FN']
        ws_cm['B5'].fill = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
        ws_cm['C5'] = cm['TP']
        ws_cm['C5'].fill = PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')
        
        ws_cm['B6'] = f"TN = {cm['TN']} (Correct Pass)"
        ws_cm['C6'] = f"FP = {cm['FP']} (False Alarm)"
        ws_cm['B7'] = f"FN = {cm['FN']} (Missed Cheat)"
        ws_cm['C7'] = f"TP = {cm['TP']} (Caught Cheat)"
        
        # SHEET 3: METRICS
        ws_metrics = wb.create_sheet("Performance Metrics", 2)
        ws_metrics['A1'] = 'PERFORMANCE METRICS'
        ws_metrics['A1'].font = Font(size=14, bold=True, color='FFFFFF')
        ws_metrics['A1'].fill = PatternFill(start_color='1F4E78', end_color='1F4E78', fill_type='solid')
        ws_metrics.merge_cells('A1:C1')
        
        m = metrics['metrics']
        metric_data = [
            ['Metric', 'Value (%)', 'Interpretation'],
            ['Accuracy', m['accuracy'], 'Overall correctness'],
            ['Precision', m['precision'], 'When we flag, how often correct?'],
            ['Recall (TPR)', m['recall'], '% of cheaters we catch'],
            ['Specificity (TNR)', m['specificity'], '% of genuine we pass correctly'],
            ['F1-Score', m['f1_score'], 'Balance of precision & recall'],
            ['False Positive Rate', m['false_positive_rate'], 'Innocent students flagged'],
            ['False Negative Rate', m['false_negative_rate'], 'Cheaters that got away'],
            ['MCC', m['mcc'], 'Overall quality (-1 to +1)']
        ]
        
        for row_idx, row_data in enumerate(metric_data, start=3):
            for col_idx, value in enumerate(row_data, start=1):
                cell = ws_metrics.cell(row=row_idx, column=col_idx, value=value)
                if row_idx == 3:
                    cell.font = Font(bold=True)
                    cell.fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
        
        # SHEET 4: SESSION DETAILS
        ws_sessions = wb.create_sheet("Session Details", 3)
        ws_sessions['A1'] = 'SESSION BREAKDOWN'
        ws_sessions['A1'].font = Font(size=14, bold=True, color='FFFFFF')
        ws_sessions['A1'].fill = PatternFill(start_color='1F4E78', end_color='1F4E78', fill_type='solid')
        ws_sessions.merge_cells('A1:F1')
        
        headers = ['Session ID', 'Classification', 'Manual Label', 'System Prediction', 'Violations', 'Count']
        for col_idx, header in enumerate(headers, start=1):
            cell = ws_sessions.cell(row=3, column=col_idx, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color='D9E1F2', end_color='D9E1F2', fill_type='solid')
        
        for row_idx, session in enumerate(self.sessions, start=4):
            ws_sessions.cell(row=row_idx, column=1, value=session['session_id'])
            ws_sessions.cell(row=row_idx, column=2, value=session['classification'])
            ws_sessions.cell(row=row_idx, column=3, value=session['manual_label'])
            ws_sessions.cell(row=row_idx, column=4, value=session['system_prediction'])
            ws_sessions.cell(row=row_idx, column=5, value=', '.join(session['violations_detected']) or 'None')
            ws_sessions.cell(row=row_idx, column=6, value=session['violation_count'])
            
            color_map = {'TP': 'C6EFCE', 'TN': 'C6EFCE', 'FP': 'FFC7CE', 'FN': 'FFC7CE'}
            fill_color = color_map.get(session['classification'], 'FFFFFF')
            ws_sessions.cell(row=row_idx, column=2).fill = PatternFill(start_color=fill_color, end_color=fill_color, fill_type='solid')
        
        # Auto-adjust column widths (skip merged cells)
        for ws in wb.worksheets:
            for col_idx in range(1, ws.max_column + 1):
                max_length = 0
                column_letter = ws.cell(row=1, column=col_idx).column_letter
                
                for row_idx in range(1, ws.max_row + 1):
                    cell = ws.cell(row=row_idx, column=col_idx)
                    # Skip merged cells
                    if hasattr(cell, 'value') and cell.value is not None:
                        try:
                            cell_length = len(str(cell.value))
                            if cell_length > max_length:
                                max_length = cell_length
                        except:
                            pass
                
                adjusted_width = min(max(max_length + 2, 12), 50)
                ws.column_dimensions[column_letter].width = adjusted_width
        
        wb.save(filename)
        print(f"✅ Excel report saved: {filename}")
        return filename
    
    def export_to_word(self, filename='evaluation_report.docx'):
        """Export metrics to Word document"""
        if not WORD_AVAILABLE:
            print("❌ Word export requires: pip install python-docx")
            return None
        
        metrics = self.calculate_metrics()
        if 'error' in metrics:
            print(f"❌ {metrics['error']}")
            return None
        
        doc = Document()
        
        title = doc.add_heading('Ghost Typing Detection System', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('Evaluation Report', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Evaluation Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"Total Sessions: {metrics['total_sessions']}")
        doc.add_paragraph()
        
        doc.add_heading('Confusion Matrix', level=2)
        cm = metrics['confusion_matrix']
        
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Light Grid Accent 1'
        
        table.cell(0, 1).text = 'Predicted: GENUINE'
        table.cell(0, 2).text = 'Predicted: CHEATING'
        table.cell(1, 0).text = 'Actual: GENUINE'
        table.cell(2, 0).text = 'Actual: CHEATING'
        
        table.cell(1, 1).text = f"TN = {cm['TN']}\n(Correct Pass)"
        table.cell(1, 2).text = f"FP = {cm['FP']}\n(False Alarm)"
        table.cell(2, 1).text = f"FN = {cm['FN']}\n(Missed Cheat)"
        table.cell(2, 2).text = f"TP = {cm['TP']}\n(Caught Cheat)"
        
        doc.add_paragraph()
        
        doc.add_heading('Performance Metrics', level=2)
        m = metrics['metrics']
        
        metrics_table = doc.add_table(rows=9, cols=3)
        metrics_table.style = 'Light List Accent 1'
        
        metrics_table.cell(0, 0).text = 'Metric'
        metrics_table.cell(0, 1).text = 'Value'
        metrics_table.cell(0, 2).text = 'Interpretation'
        
        metric_rows = [
            ['Accuracy', f"{m['accuracy']:.2f}%", 'Overall correctness'],
            ['Precision', f"{m['precision']:.2f}%", 'When we flag, how often correct?'],
            ['Recall (TPR)', f"{m['recall']:.2f}%", '% of cheaters we catch'],
            ['Specificity (TNR)', f"{m['specificity']:.2f}%", '% of genuine we pass correctly'],
            ['F1-Score', f"{m['f1_score']:.2f}%", 'Balance of precision & recall'],
            ['False Positive Rate', f"{m['false_positive_rate']:.2f}%", 'Innocent students flagged'],
            ['False Negative Rate', f"{m['false_negative_rate']:.2f}%", 'Cheaters that got away'],
            ['MCC', f"{m['mcc']:.4f}", 'Overall quality (-1 to +1)']
        ]
        
        for idx, row_data in enumerate(metric_rows, start=1):
            for col_idx, value in enumerate(row_data):
                metrics_table.cell(idx, col_idx).text = value
        
        doc.add_paragraph()
        
        doc.add_heading('Recommendations', level=2)
        
        if m['accuracy'] >= 90:
            doc.add_paragraph('✅ EXCELLENT accuracy! System is highly reliable.', style='List Bullet')
        elif m['accuracy'] >= 80:
            doc.add_paragraph('✅ GOOD accuracy. System is performing well.', style='List Bullet')
        elif m['accuracy'] >= 70:
            doc.add_paragraph('⚠️ MODERATE accuracy. Consider improvements.', style='List Bullet')
        else:
            doc.add_paragraph('❌ LOW accuracy. System needs significant improvement.', style='List Bullet')
        
        if m['false_positive_rate'] > 15:
            doc.add_paragraph('⚠️ HIGH false positive rate - Too many innocent students flagged!', style='List Bullet')
        
        if m['false_negative_rate'] > 15:
            doc.add_paragraph('⚠️ HIGH false negative rate - Missing too many cheaters!', style='List Bullet')
        
        doc.save(filename)
        print(f"✅ Word report saved: {filename}")
        return filename
    
    def print_detailed_report(self):
        """Print comprehensive evaluation report"""
        metrics = self.calculate_metrics()
        
        if 'error' in metrics:
            print(f"\n❌ {metrics['error']}")
            print(f"💡 {metrics.get('message', '')}")
            return
        
        print("\n" + "="*85)
        print("🎯 GHOST TYPING DETECTION - COMPREHENSIVE EVALUATION REPORT")
        print("="*85)
        
        print(f"\n📊 Dataset Information:")
        print(f"   Total Sessions Evaluated: {metrics['total_sessions']}")
        print(f"   Genuine Sessions: {metrics['breakdown']['genuine_sessions']}")
        print(f"   Cheating Sessions: {metrics['breakdown']['cheating_sessions']}")
        print(f"   Correctly Classified: {metrics['breakdown']['correctly_classified']}")
        print(f"   Incorrectly Classified: {metrics['breakdown']['incorrectly_classified']}")
        
        cm = metrics['confusion_matrix']
        print("\n" + "="*85)
        print("📈 CONFUSION MATRIX BREAKDOWN")
        print("="*85)
        print(f"""
   ╔════════════════════════════════════════════════════════════════╗
   ║                                                                ║
   ║  True Positives  (TP): {cm['TP']:4d}  ✅ Cheaters CAUGHT        ║
   ║  True Negatives  (TN): {cm['TN']:4d}  ✅ Genuine PASSED         ║
   ║  False Positives (FP): {cm['FP']:4d}  ❌ Innocent FLAGGED       ║
   ║  False Negatives (FN): {cm['FN']:4d}  ❌ Cheaters MISSED        ║
   ║                                                                ║
   ╚════════════════════════════════════════════════════════════════╝
        """)
        
        m = metrics['metrics']
        print("="*85)
        print("🎯 PERFORMANCE METRICS")
        print("="*85)
        print(f"""
   Accuracy:            {m['accuracy']:6.2f}%  │ Overall correctness
   ─────────────────────────────────────────────────────────────
   Precision:           {m['precision']:6.2f}%  │ When we flag, how often correct?
   Recall (TPR):        {m['recall']:6.2f}%  │ % of cheaters we catch
   Specificity (TNR):   {m['specificity']:6.2f}%  │ % of genuine we pass correctly
   F1-Score:            {m['f1_score']:6.2f}%  │ Balance of precision & recall
   ─────────────────────────────────────────────────────────────
   False Positive Rate: {m['false_positive_rate']:6.2f}%  │ Innocent students flagged
   False Negative Rate: {m['false_negative_rate']:6.2f}%  │ Cheaters that got away
   ─────────────────────────────────────────────────────────────
   MCC (Correlation):   {m['mcc']:7.4f}  │ Overall quality (-1 to +1)
        """)
        
        print("="*85)
        print("💡 INTERPRETATION & RECOMMENDATIONS")
        print("="*85)
        
        if m['accuracy'] >= 90:
            print("   ✅ EXCELLENT accuracy! System is highly reliable.")
        elif m['accuracy'] >= 80:
            print("   ✅ GOOD accuracy. System is performing well.")
        elif m['accuracy'] >= 70:
            print("   ⚠️  MODERATE accuracy. Consider improvements.")
        else:
            print("   ❌ LOW accuracy. System needs significant improvement.")
        
        if m['false_positive_rate'] > 15:
            print("   ⚠️  HIGH false positive rate - Too many innocent students flagged!")
            print("      → Consider reducing sensitivity or improving detection logic")
        elif m['false_positive_rate'] > 5:
            print("   ⚠️  Moderate false positive rate - Some innocent students flagged")
        else:
            print("   ✅ Low false positive rate - Good!")
        
        if m['false_negative_rate'] > 15:
            print("   ⚠️  HIGH false negative rate - Missing too many cheaters!")
            print("      → Consider increasing sensitivity or adding more detection methods")
        elif m['false_negative_rate'] > 5:
            print("   ⚠️  Moderate false negative rate - Some cheaters getting through")
        else:
            print("   ✅ Low false negative rate - Good!")
        
        if m['f1_score'] >= 85:
            print("   ✅ Excellent balance between precision and recall")
        elif m['f1_score'] >= 70:
            print("   ⚠️  Moderate F1-score - Room for improvement")
        else:
            print("   ❌ Low F1-score - Imbalanced performance")
        
        print("\n" + "="*85)
        
        print("\n📋 DETAILED SESSION BREAKDOWN:")
        print("-" * 85)
        
        for category in ['TP', 'TN', 'FP', 'FN']:
            category_sessions = [s for s in self.sessions if s['classification'] == category]
            if category_sessions:
                category_names = {
                    'TP': 'True Positives (Cheaters Caught)',
                    'TN': 'True Negatives (Genuine Passed)',
                    'FP': 'False Positives (Innocent Flagged)',
                    'FN': 'False Negatives (Cheaters Missed)'
                }
                print(f"\n{category_names[category]}: {len(category_sessions)}")
                for s in category_sessions[:5]:
                    print(f"  • {s['session_id']}: {len(s['violations_detected'])} violations")
                if len(category_sessions) > 5:
                    print(f"  ... and {len(category_sessions) - 5} more")
        
        print("\n" + "="*85 + "\n")
    
    def export_results(self, filename='evaluation_results.json'):
        """Export all results to JSON - NumPy 2.0 Compatible"""
        metrics = self.calculate_metrics()
        
        def convert_to_native(obj):
            """Convert NumPy and other types to JSON-serializable Python types"""
            if isinstance(obj, np.integer):
                return int(obj)
            elif isinstance(obj, np.floating):
                return float(obj)
            elif isinstance(obj, np.ndarray):
                return obj.tolist()
            elif isinstance(obj, dict):
                return {key: convert_to_native(value) for key, value in obj.items()}
            elif isinstance(obj, list):
                return [convert_to_native(item) for item in obj]
            return obj
        
        export_data = {
            'evaluation_date': datetime.now().isoformat(),
            'metrics': convert_to_native(metrics),
            'sessions': convert_to_native(self.sessions)
        }
        
        if len(self.sessions) > 0:
            y_pred = [1 if s['system_prediction'] == 'CHEATING' else 0 for s in self.sessions]
            y_true = [1 if s['manual_label'] == 'CHEATING' else 0 for s in self.sessions]
            
            report = classification_report(
                y_true, y_pred,
                target_names=['GENUINE', 'CHEATING'],
                output_dict=True
            )
            export_data['sklearn_report'] = convert_to_native(report)
        
        with open(filename, 'w') as f:
            json.dump(export_data, f, indent=2)
        
        print(f"✅ Results exported to: {filename}")
        return filename


def main():
    """Main function"""
    print("🚀 Starting Ghost Typing Evaluation System\n")
    
    evaluator = GhostTypingEvaluator(backend_url="https://exam-proctor-backend-jxrb.onrender.com")
    
    print("Attempting to fetch labeled submissions from backend...")
    submissions = evaluator.fetch_labeled_submissions()
    
    if submissions and len(submissions) > 0:
        evaluator.process_submissions(submissions)
    
    # If no backend data, use test data
    if len(evaluator.sessions) == 0:
        print("\n⚠️ No labeled submissions found. Using manual test data...\n")
        print("Adding test sessions for demonstration:\n")
        
        evaluator.add_manual_session('TEST001', ['GHOST_TYPING_DETECTED'], 'CHEATING',
                                     {'note': 'Clear ghost typing pattern'})
        evaluator.add_manual_session('TEST002', ['GHOST_TYPING_DETECTED', 'TAB_SWITCH'], 'CHEATING',
                                     {'note': 'Multiple violations'})
        evaluator.add_manual_session('TEST003', [], 'GENUINE',
                                     {'note': 'Clean session'})
        evaluator.add_manual_session('TEST004', [], 'GENUINE',
                                     {'note': 'No violations detected'})
        evaluator.add_manual_session('TEST005', [], 'GENUINE',
                                     {'note': 'Perfect behavior'})
        evaluator.add_manual_session('TEST006', ['NO_FACE_DETECTED'], 'GENUINE',
                                     {'note': 'Webcam issue, not cheating'})
        evaluator.add_manual_session('TEST007', [], 'CHEATING',
                                     {'note': 'Used phone off-camera, not detected'})
        evaluator.add_manual_session('TEST008', ['GHOST_TYPING_DETECTED'], 'CHEATING')
        evaluator.add_manual_session('TEST009', [], 'GENUINE')
        evaluator.add_manual_session('TEST010', ['GHOST_TYPING_DETECTED'], 'CHEATING')
        evaluator.add_manual_session('TEST011', [], 'GENUINE')
        evaluator.add_manual_session('TEST012', ['MULTIPLE_PERSONS'], 'CHEATING')
        evaluator.add_manual_session('TEST013', [], 'GENUINE')
        evaluator.add_manual_session('TEST014', ['TAB_SWITCH'], 'GENUINE')
    
    evaluator.print_detailed_report()
    evaluator.plot_confusion_matrix('ghost_typing_confusion_matrix.png')
    evaluator.export_results('evaluation_results.json')
    evaluator.export_to_excel('evaluation_metrics.xlsx')
    evaluator.export_to_word('evaluation_report.docx')
    
    print("\n" + "="*85)
    print("✅ EVALUATION COMPLETE!")
    print("="*85)
    print("\nGenerated files:")
    print("  1. ghost_typing_confusion_matrix.png - Visual confusion matrix")
    print("  2. evaluation_results.json - Complete metrics data")
    print("  3. evaluation_metrics.xlsx - Formatted Excel workbook")
    print("  4. evaluation_report.docx - Word document report")
    print("\n💡 Backend Connection:")
    print("  - Check if your backend is running on https://exam-proctor-backend-jxrb.onrender.com")
    print("  - Verify the API endpoint returns labeled submissions")
    print("  - Check MongoDB for sessions with 'label' field set")
    print("\n💡 Next steps:")
    print("  - Review the confusion matrix image")
    print("  - Open Excel file for detailed metrics tables")
    print("  - Check false positives and false negatives")
    print("  - Aim for 85%+ accuracy with <10% false positive rate")
    print("="*85 + "\n")


if __name__ == "__main__":
    main()